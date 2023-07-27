from re import T
from tkinter import *
import os
from PIL import ImageTk, Image
import moviepy.editor as mp
from docx2pdf import convert
from pdf2docx import parse
from fpdf import FPDF
from docx import Document
import aspose.words as aw
import docx2txt
import subprocess
from moviepy.editor import *
import ffmpy
from tkinter.filedialog import *
import re

# creating root object
Screen = Tk()
Screen.geometry("1920x1080")
Screen.title("Ani-vertor | An Allrounder Convertor.")
Screen.config(bg='#A3BBAD')

# Creating frames
Tops = Frame(Screen, width=2600, relief=SUNKEN,  bg='#A3BBAD')
Tops.pack(side=TOP)
f1 = Frame(Screen, width=800, height=700, relief=SUNKEN, bg='#A3BBAD')
f1.pack(side=BOTTOM)
addr = StringVar()
Result = StringVar()
clickedfrom = StringVar()
clickedto = StringVar()

#  Function to browse file and store it's into a variable as string
def OpenFile():
    global FileOpen
    FileOpen = StringVar()
    FileOpen = askopenfilename(initialdir="/Desktop", title="SelectFile",
                               filetypes=(("only jpeg files", "jpg"), ("all type of files", ".*")))
    addr.set(FileOpen)


def runCode():
    from_type = clickedfrom.get()
    to_type = clickedto.get()
    file_path = addr.get()
    file_path = file_path.replace("/", "\\\\")
    img_folder_path = "C:\\Ani-vertor\\images"
    video_folder_path = "C:\\Ani-vertor\\videos"
    txt_documents_folder_path = "C:\\Ani-vertor\\documents\\TextFile"
    pdf_documents_folder_path = "C:\\Ani-vertor\\documents\\PDFFile"
    word_documents_folder_path = "C:\\Ani-vertor\\documents\\WordFile"
    audio_folder_path = "C:\\Ani-vertor\\audios"
    gif_folder_path = "C:\\Ani-vertor\\gifs"

    if from_type == to_type:
        Result.set("Select different file types!")
        return
    else:
        extensions = {
            "jpg": 1,
            "png": 2,
            "jpeg": 3,
            "ico": 4,
            "txt": 5,
            "pdf": 6,
            "docx": 7,
            "mp3": 8,
            "wav": 9,
            "mp4": 10,
            "gif": 11,
        }

    def switcher(argument):
        match argument:
            case 1:
                if to_type == "png": # jpg to png
                    os.chdir(img_folder_path)
                    for root, dirs, files in os.walk(img_folder_path):
                        for file_name in files:
                            img = Image.open(os.path.join(root, file_name))

                    tempResult = file_name.split('.', 1)[0] + "_copy." + to_type

                    img.save(tempResult)
                    img.close()
                    Result.set("File converted Successfully !!")
                    return 0
                
                elif to_type == "jpeg": # jpg to jpeg
                    os.chdir(img_folder_path)
                    for root, dirs, files in os.walk(img_folder_path):
                        for file_name in files:
                            img = Image.open(os.path.join(root, file_name))

                    tempResult = file_name.split('.', 1)[0] + "_copy." + to_type

                    img.save(tempResult)
                    img.close()
                    Result.set("File converted Successfully !!")
                    return 0
                
                elif to_type == "ico": # jpg to ico
                    os.chdir(img_folder_path)
                    for root, dirs, files in os.walk(img_folder_path):
                        for file_name in files:
                            img = Image.open(os.path.join(root, file_name))

                    tempResult = file_name.split('.', 1)[0] + "_copy." + to_type

                    img.save(tempResult)
                    img.close()
                    Result.set("File converted Successfully !!")
                    return 0

            case 2:
                if to_type == "jpg": # png to jpg
                    os.chdir(img_folder_path)
                    for root, dirs, files in os.walk(img_folder_path):
                        for file_name in files:
                            img = Image.open(os.path.join(root, file_name))

                    tempResult = file_name.split('.', 1)[0] + "_copy." + to_type

                    img.save(tempResult)
                    img.close()
                    Result.set("File converted Successfully !!")
                    return 0
                
                elif to_type == "jpeg": # png to jpeg
                    os.chdir(img_folder_path)
                    for root, dirs, files in os.walk(img_folder_path):
                        for file_name in files:
                            img = Image.open(os.path.join(root, file_name))

                    tempResult = file_name.split('.', 1)[0] + "_copy." + to_type

                    img.save(tempResult)
                    img.close()
                    Result.set("File converted Successfully !!")
                    return 0
                elif to_type == "ico": # png to ico
                    os.chdir(img_folder_path)
                    for root, dirs, files in os.walk(img_folder_path):
                        for file_name in files:
                            img = Image.open(os.path.join(root, file_name))

                    tempResult = file_name.split('.', 1)[0] + "_copy." + to_type

                    img.save(tempResult)
                    img.close()
                    Result.set("File converted Successfully !!")
                    return 0

            case 3:
                if to_type == "jpg": # jpeg to jpg
                    os.chdir(img_folder_path)
                    for root, dirs, files in os.walk(img_folder_path):
                        for file_name in files:
                            img = Image.open(os.path.join(root, file_name))

                    tempResult = file_name.split('.', 1)[0] + "_copy." + to_type

                    img.save(tempResult)
                    img.close()
                    Result.set("File converted Successfully !!")
                    return 0
                
                elif to_type == "png": # jpeg to png
                    os.chdir(img_folder_path)
                    for root, dirs, files in os.walk(img_folder_path):
                        for file_name in files:
                            img = Image.open(os.path.join(root, file_name))

                    tempResult = file_name.split('.', 1)[0] + "_copy." + to_type

                    img.save(tempResult)
                    img.close()
                    Result.set("File converted Successfully !!")
                    return 0
                
                elif to_type == "ico": # jpeg to ico
                    os.chdir(img_folder_path)
                    for root, dirs, files in os.walk(img_folder_path):
                        for file_name in files:
                            img = Image.open(os.path.join(root, file_name))

                    tempResult = file_name.split('.', 1)[0] + "_copy." + to_type

                    img.save(tempResult)
                    img.close()
                    Result.set("File converted Successfully !!")
                    return 0

            case 4:
                if to_type == "jpg": # ico to jpg
                    os.chdir(img_folder_path)
                    for root, dirs, files in os.walk(img_folder_path):
                        for file_name in files:
                            img = Image.open(os.path.join(root, file_name))

                    tempResult = file_name.split('.', 1)[0] + "_copy." + to_type

                    img.save(tempResult)
                    img.close()
                    Result.set("File converted Successfully !!")
                    return 0
                
                elif to_type == "png": # ico to png
                    os.chdir(img_folder_path)
                    for root, dirs, files in os.walk(img_folder_path):
                        for file_name in files:
                            img = Image.open(os.path.join(root, file_name))

                    tempResult = file_name.split('.', 1)[0] + "_copy." + to_type

                    img.save(tempResult)
                    img.close()
                    Result.set("File converted Successfully !!")
                    return 0
                
                elif to_type == "jpeg": # ico to jpeg
                    os.chdir(img_folder_path)
                    for root, dirs, files in os.walk(img_folder_path):
                        for file_name in files:
                            img = Image.open(os.path.join(root, file_name))

                    tempResult = file_name.split('.', 1)[0] + "_copy." + to_type

                    img.save(tempResult)
                    img.close()
                    Result.set("File converted Successfully !!")
                    return 0

            case 5: 
                if to_type == "pdf": # txt to pdf
                    os.chdir(txt_documents_folder_path)
                    pdf = FPDF()  # save FPDF() class into a variable pdf
                    pdf.add_page()  # Add a page
                    pdf.set_font("Arial", size=15)
                    for root, dirs, files in os.walk(txt_documents_folder_path):
                        for file_name in files:
                            f = open(os.path.join(root, file_name))

                    for x in f:  # insert the texts in pdf
                        pdf.cell(200, 10, txt=x, ln=1, align='L')

                    tempResult = file_name.split('.', 1)[0] + "_copy." + to_type
                    pdf.output(tempResult)  # save the pdf with name .pdf
                    Result.set("File converted Successfully !!")
                    return 0
                
                elif to_type == "docx": # txt to docs
                    direct = os.listdir(txt_documents_folder_path)
                    
                    for i in direct:
                        document = Document()
                        document.add_heading(i, 0)
                        myfile = open( txt_documents_folder_path +"\\"+i).read()
                        myfile = re.sub(r'[^\x00-\x7F]+|\x0c',' ', myfile) # remove all non-XML-compatible characters
                        p = document.add_paragraph(myfile)
                        tempResult = i.split('.', 1)[0] + "_copy." + to_type
                        document.save(txt_documents_folder_path + "\\" + tempResult)
                        Result.set("File converted Successfully !!")
                    return 0
                
            case 6:
                if to_type == "txt": # pdf to txt
                    os.chdir(pdf_documents_folder_path)
                    for root, dirs, files in os.walk(pdf_documents_folder_path):
                        for file_name in files:
                            doc = aw.Document(os.path.join(root, file_name))
                    
                    tempResult = file_name.split('.', 1)[0] + "_copy." + to_type
                    doc.save(tempResult)
                    
                    Result.set("File converted Successfully !!")
                    return 0

                elif to_type == "docx": # pdf to docs
                    os.chdir(pdf_documents_folder_path)
                    for root, dirs, files in os.walk(pdf_documents_folder_path):
                        for file_name in files:
                            doc = aw.Document(os.path.join(root, file_name))
                    parse(pdf_documents_folder_path+"\\"+file_name,
                        pdf_documents_folder_path +"\\" + file_name.split('.', 1)[0] + "_copy." + to_type)
                    Result.set("File converted Successfully !!")
                    return 0
                
            case 7:
                if to_type == "txt": # docs to txt
                    os.chdir(word_documents_folder_path)
                    for root, dirs, files in os.walk(word_documents_folder_path):
                        for file_name in files:
                            doc = aw.Document(os.path.join(root, file_name))
                    
                    MY_TEXT = docx2txt.process(file_name)
                    with open(word_documents_folder_path+"\\"+file_name.split('.', 1)[0] + "_copy." + to_type, "w") as text_file:
                        print(MY_TEXT, file=text_file)
                    Result.set("File converted Successfully !!")
                    return 0
                
                elif to_type == "pdf": # docs to pdf
                    os.chdir(word_documents_folder_path)
                    for root, dirs, files in os.walk(word_documents_folder_path):
                        for file_name in files:
                            doc = aw.Document(os.path.join(root, file_name))
                    
                    file_name = word_documents_folder_path +"\\"+file_name
                    tempResult = file_name.split('.', 1)[0] + "_copy." + to_type
                    convert(file_name,tempResult)
                    Result.set("File converted Successfully !!")
                    return 0
                
            case 8:
                if to_type == "wav": # mp3 to wav
                    os.chdir(audio_folder_path)
                    for root, dirs, files in os.walk(audio_folder_path):
                        for file_name in files:
                            doc = os.path.join(root, file_name)
                    
                    file_name = audio_folder_path +"\\"+ file_name
                    tempResult = file_name.split('.', 1)[0] + "_copy." + to_type
                    subprocess.call(['ffmpeg', '-i', file_name,tempResult])                    
                    Result.set("File converted Successfully !!")
                    return 0
                
            case 9: # wav to mp3
                if to_type == "mp3":
                    Result.set("Choose different file format")
                    return 0
                
            case 10:
                if to_type == "mp3": # mp4 to mp3
                    os.chdir(video_folder_path)
                    for root, dirs, files in os.walk(video_folder_path):
                        for file_name in files:
                            doc = os.path.join(root, file_name)

                    file_name = video_folder_path +"\\"+ file_name
                    tempResult = file_name.split('.', 1)[0] + "_copy." + to_type

                    clip = mp.VideoFileClip(file_name)
                    clip.audio.write_audiofile(tempResult)
                    Result.set("File converted Successfully !!")
                    return 0
                
 		    elif to_type == "gif": # mp4 to gif
                    os.chdir(video_folder_path)
                    for root, dirs, files in os.walk(video_folder_path):
                        for file_name in files:
                            doc = os.path.join(root, file_name)

                    file_name = video_folder_path +"\\"+ file_name
                    tempResult = file_name.split('.', 1)[0] + "_copy." + to_type

                    clip = (VideoFileClip(file_name))
                    clip.write_gif(tempResult)
                    Result.set("File converted Successfully !!")
                    return 0

            case 11:
                if to_type == "mp4": # gif to mp4
                    os.chdir(gif_folder_path)
                    for root, dirs, files in os.walk(gif_folder_path):
                        for file_name in files:
                            doc = os.path.join(root, file_name)

                    file_name = gif_folder_path +"\\"+ file_name
                    tempResult = file_name.split('.', 1)[0] + "_copy." + to_type

                    ff = ffmpy.FFmpeg(
                        inputs={file_name: None},
                        outputs={tempResult: None}
                    )
                    ff.run()
                    Result.set("File converted Successfully.")
                    return 0
                
                elif to_type == "mp3": # gif to mp3
                    Result.set("Cannot Perform Task.")
                    return 0
                
            case default:
                Result.set("Choose different file formats.")
                return 0

    switcher(extensions[from_type])

# Column for Heading
lblInfo = Label(Tops, font=('helvetica', 45, 'bold'),
                bg='#A3BBAD', text="", fg="Black", bd=10, anchor='w')
lblInfo.grid(row=0, column=0)
lblInfo = Label(Tops, font=('helvetica', 45, 'bold'),
                text="Ani-vertor", fg="Black", bd=10, anchor='w')
lblInfo.grid(row=1, column=0)

def showfrom():
    labelfrom.config(text=clickedfrom.get())

fromoptions = [
    "jpg",
    "png",
    "jpeg",
    "ico",
    "txt",
    "pdf",
    "docx",
    "mp3",
    "wav",
    "mp4",
    "gif"
]

clickedfrom.set("jpg")
dropfrom = OptionMenu(Screen, clickedfrom, *fromoptions)
dropfrom.place(relx=0.1, rely=0.5)
dropfrom.configure(bg='#00A6ED')

fromBtn = Button(f1, padx=5, pady=3, bd=5, fg="black",
                 font=('arial', 16, 'bold'), width=10,
                 text="click Me", bg="#A8B4A5",
                 command=showfrom).grid(row=0, column=1)

labelfrom = Label(f1, font=('arial', 10, 'bold'), text=" ", bd=5, anchor="w")
labelfrom.grid(row=0, column=2)

def showto():
    labelto.config(text=clickedto.get())

tooptions = [
    "jpg",
    "png",
    "jpeg",
    "ico",
    "txt",
    "pdf",
    "docx",
    "mp3",
    "wav",
    "mp4",
    "gif"
]
clickedto.set("jpg")
dropto = OptionMenu(Screen, clickedto, *tooptions)
dropto.place(relx=0.1, rely=0.76)
dropto.configure(bg='#00A6ED')

toBtn = Button(f1, padx=5, pady=3, bd=5, fg="black",
               font=('arial', 16, 'bold'), width=10,
               text="click Me", bg="#A8B4A5",
               command=showto).grid(row=3, column=1)

labelto = Label(f1, font=('arial', 10, 'bold'), text=" ", bd=5, anchor="w")
labelto.grid(row=3, column=2)

searchFile = Button(f1, padx=5, pady=3, bd=5, fg="black",
                    font=('arial', 16, 'bold'), width=10,
                    text="Search file", bg="#00A6ED",
                    command=OpenFile).grid(row=1, column=0)

# 2 Empty columns for blank spaces
emptylabel = Label(f1, font=('arial', 16, 'bold'),
                   text="", width=10,  bd=16, anchor="w")
emptylabel.grid(row=0, column=3)
emptylabel.configure(bg='#A3BBAD')
emptylabel1 = Label(f1, font=('arial', 16, 'bold'), text="", bd=16, anchor="w")
emptylabel1.grid(row=0, column=4)
emptylabel1.configure(bg='#A3BBAD')

# Column for file name display
addrReference = Entry(f1, font=('arial', 16, 'bold'),
                      textvariable=addr, justify='left')
addrReference.grid(row=1, column=1)

# Empty column for blank spaces
emptylabel2 = Label(f1, font=('arial', 16, 'bold'), text="", bd=16, anchor="w")
emptylabel2.grid(row=2, column=0)
emptylabel2.configure(bg='#A3BBAD')
emptylabel3 = Label(f1, font=('arial', 16, 'bold'),
                    text="", width=10,  bd=16, anchor="w")
emptylabel3.grid(row=3, column=0)
emptylabel3.configure(bg='#A3BBAD')

processIntr = Button(f1, padx=5, pady=3, bd=5, fg="black",
                     font=('arial', 16, 'bold'), width=10,
                     text="Convert", bg="#A8B4A5",
                     command=runCode).grid(row=0, column=6)

# Column for show result
showResult = Label(f1, font=('arial', 16, 'bold'),
                   text="Result :", bd=16, anchor="w")
showResult.grid(row=4,  column=4)
showResult.configure(bg='#A3BBAD')

# Column for show result output
resultMessage = Entry(f1, font=('arial', 16, 'bold'), textvariable=Result,
                      bd=10, insertwidth=4, width=30, bg="white", justify='left')
resultMessage.grid(row=4, columnspan=2, column=5)

# Empty column for blank spaces
emptylabel2 = Label(f1, font=('arial', 16, 'bold'), text="", bd=16, anchor="w")
emptylabel2.grid(row=5, column=0)
emptylabel2.configure(bg='#A3BBAD')

Screen.mainloop()