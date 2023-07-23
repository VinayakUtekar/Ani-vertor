from re import T
from tkinter import *

# import other necessery modules
from PIL import ImageTk,Image
import moviepy.editor as mp
from docx2pdf import convert
from pdf2docx import parse
from fpdf import FPDF
from docx import Document
import glob
import os
import aspose.words as aw
import docx2txt
import subprocess 
from moviepy.editor import *
import ffmpy
from tkinter.filedialog import *


# creating root object
Screen = Tk()
Screen.geometry("1920x1080")
Screen.title("Ani-vertor | An Allrounder Convertor.")
Screen.config(bg='#A3BBAD')

# Creating frames 
Tops = Frame(Screen, width = 2600, relief = SUNKEN,  bg='#A3BBAD')
Tops.pack(side = TOP)
f1 = Frame(Screen, width = 800, height = 700, relief = SUNKEN, bg='#A3BBAD')
f1.pack(side = BOTTOM)
addr = StringVar()
Result = StringVar()
clickedfrom = StringVar()
clickedto = StringVar()



#  Function to browse file and store it's into a variable as string
def OpenFile():
    global FileOpen
    FileOpen=StringVar()
    FileOpen = askopenfilename(initialdir ="/Desktop",title="SelectFile",filetypes=(("only jpeg files","jpg"),("all type of files",".*")))
    addr.set(FileOpen)


def runCode():
    from_type = clickedfrom.get()
    to_type = clickedto.get()
    file_path = addr.get()
    # print((file_path))
    if from_type == to_type:
        Result.set("Select different file types!")
        return 
    else:
        extensions = {
            "jpg" : 1,
            "png" : 2,
            "jpeg" : 3,
            "ico" : 4,
            "txt" : 5,
            "pdf" : 6,
            "docs" : 7,
            "mp3" : 8,
            "wav" : 9,
            "mp4" : 10,
            "gif" : 11,
        }
    def switcher(argument):
        match argument:
            case 1:
                if to_type == "png":
                    im1 = Image.open(r'C:\Ani-vertor\pic.jpg')
                    im1.save(r'C:\Ani-vertor\pic.png')
                    im1.close()
                    Result.set("File converted Successfully !!")
                    return 0
                elif to_type == "jpeg":
                    im1 = Image.open(r'C:\Ani-vertor\pic.jpg')
                    im1.save(r'C:\Ani-vertor\pic.jpeg')
                    im1.close()
                    Result.set("File converted Successfully !!")
                    return 0
                elif to_type == "ico":
                    im1 = Image.open(r'C:\Ani-vertor\pic.jpg')
                    im1.save(r'C:\Ani-vertor\pic.ico')
                    im1.close()
                    Result.set("File converted Successfully !!")
                    return 0

            case 2:
                if to_type == "jpg":
                    im1 = Image.open(r'C:\Ani-vertor\pic.png')
                    im1.save(r'C:\Ani-vertor\pic.jpg')
                    im1.close()
                    Result.set("File converted Successfully !!")
                    return 0
                elif to_type == "jpeg":
                    im1 = Image.open(r'C:\Ani-vertor\pic.png')
                    im1.save(r'C:\Ani-vertor\pic.jpeg')
                    im1.close()
                    Result.set("File converted Successfully !!")
                    return 0
                elif to_type == "ico":
                    im1 = Image.open(r'C:\Ani-vertor\pic.png')
                    im1.save(r'C:\Ani-vertor\pic.ico')
                    im1.close()
                    Result.set("File converted Successfully !!")
                    return 0
            
            case 3:
                if to_type == "jpg":
                    im1 = Image.open(r'C:\Ani-vertor\pic.jpeg')
                    im1.save(r'C:\Ani-vertor\pic.jpg')
                    im1.close()
                    Result.set("File converted Successfully !!")
                    return 0
                elif to_type == "jpeg":
                    im1 = Image.open(r'C:\Ani-vertor\pic.jpeg')
                    im1.save(r'C:\Ani-vertor\pic.png')
                    im1.close()
                    Result.set("File converted Successfully !!")
                    return 0
                elif to_type == "ico":
                    im1 = Image.open(r'C:\\Ani-vertor\\pic.jpeg')
                    im1.save(r'C:\Ani-vertor\pic.ico')
                    im1.close()
                    Result.set("File converted Successfully !!")
                    return 0
            
            case 4:
                if to_type == "jpg":
                    im1 = Image.open(r'C:\Ani-vertor\pic.ico')
                    im1.save(r'C:\Ani-vertor\pic.jpg')
                    im1.close()
                    Result.set("File converted Successfully !!")
                    return 0
                elif to_type == "png":
                    im1 = Image.open(r'C:\Ani-vertor\pic.ico')
                    im1.save(r'C:\Ani-vertor\pic.png')
                    im1.close()
                    Result.set("File converted Successfully !!")
                    return 0
                elif to_type == "jpeg":
                    im1 = Image.open(r'C:\Ani-vertor\pic.ico')
                    im1.save(r'C:\Ani-vertor\pic.jpeg')
                    im1.close()
                    Result.set("File converted Successfully !!")
                    return 0

            case 5:
                if to_type == "pdf":
                    pdf = FPDF() # save FPDF() class into a variable pdf
                    pdf.add_page() # Add a page
                    pdf.set_font("Arial", size = 15) # set style and size of font that you want in the pdf
                    f = open("C:\Ani-vertor\SampleTextFile.txt", "r") # open the text file in read mode
                    for x in f: # insert the texts in pdf
                        pdf.cell(200, 10, txt = x, ln = 1, align = 'L')
                    pdf.output("C:\Ani-vertor\SampleTextFile.pdf") # save the pdf with name .pdf
                    Result.set("File converted Successfully !!")
                    return 0
                elif to_type == "docs":
                    doc = Document()
                    files = glob.glob("*.txt")
                    print(files)
                    file = input("C:\Ani-vertor\SampleTextFile.txt")
                    
                    with open(file, 'r', encoding='utf-8') as openfile:
                        line = openfile.read()
                        doc.add_paragraph(line)
                        doc.save(file + ".docx")
                    
                    os.system(file + ".docx")
                    Result.set("File converted Successfully !!")
                    return 0
            case 6:
                if to_type == "txt":
                    doc = aw.Document("C:\Ani-vertor\SamplePDFFile.pdf")
                    doc.save("C:\Ani-vertor\SamplePDFFile.txt")
                    Result.set("File converted Successfully !!")
                    return 0
                elif to_type == "docs":
                    parse(r"C:\Ani-vertor\SamplePDFFile.pdf",r"C:\Ani-vertor\SamplePDFFile.docx")
                    Result.set("File converted Successfully !!")
                    return 0
            case 7:
                if to_type == "txt":
                    MY_TEXT = docx2txt.process("C:\Ani-vertor\SampleWordFile.docx")
                    with open("C:\Ani-vertor\SampleWordFile.txt", "w") as text_file:
                        print(MY_TEXT, file=text_file)
                    Result.set("File converted Successfully !!")
                    return 0
                elif to_type == "pdf":
                    convert(r"C:\Ani-vertor\SampleWordFile.docx",r"C:\Ani-vertor\SampleWordFile.pdf")
                    Result.set("File converted Successfully !!")
                    return 0
            case 8:
                if to_type == "wav":
                    # convert mp3 to wav file 
                    subprocess.call(['ffmpeg', '-i', 'C:\Ani-vertor\CountOnMe.mp3', 'C:\Ani-vertor\CountOnMe.wav'])
                    Result.set("File converted Successfully !!")
                    return 0
            case 9:
                if to_type == "":

                    return 0
            case 10:
                if to_type == "mp3":
                    clip = mp.VideoFileClip(r"C:\Ani-vertor\Video.mp4")
                    clip.audio.write_audiofile(r"C:\Ani-vertor\Video.mp3")
                    Result.set("File converted Successfully !!")
                    return 0
                elif to_type == "gif":
                    clip = (VideoFileClip("C:\Ani-vertor\Video.mp4"))
                    clip.write_gif("C:\Ani-vertor\Video.gif")
                    # replace following line with location of your .docx file

            case 11:
                
                ff = ffmpy.FFmpeg(
                    inputs={'C:\Ani-vertor\giphy.gif': None},
                    outputs={'C:\Ani-vertor\giphy.mp4': None}
                )
                ff.run()
                return 0
            case default:
                return 0


    switcher(extensions[from_type])
                







# Column for Heading
lblInfo = Label(Tops, font = ('helvetica', 45, 'bold'), bg='#A3BBAD', text = "", fg = "Black", bd = 10, anchor='w')
lblInfo.grid(row = 0, column = 0)
lblInfo = Label(Tops, font = ('helvetica', 45, 'bold'), text = "Ani-vertor", fg = "Black", bd = 10, anchor='w')
lblInfo.grid(row = 1, column = 0)



def showfrom():
    labelfrom.config( text = clickedfrom.get() )
fromoptions = [
    "jpg",
    "png",
    "jpeg",
    "ico",
    "txt",
    "pdf",
    "docs",
    "mp3",
    "wav",
    "mp4",
    "gif"
]

clickedfrom.set( "jpg" )
dropfrom = OptionMenu( Screen , clickedfrom , *fromoptions )
dropfrom.place(relx=0.1,rely=0.5)
dropfrom.configure(bg='#00A6ED')

fromBtn = Button(f1, padx = 5, pady = 3, bd = 5, fg = "black",
font = ('arial', 16, 'bold'), width = 10,
text = "click Me", bg = "#A8B4A5",
command = showfrom).grid(row = 0, column = 1)

labelfrom = Label( f1 , font = ('arial', 10, 'bold'), text = " ", bd = 5, anchor = "w" )
labelfrom.grid(row = 0, column = 2)






def showto():
    labelto.config( text = clickedto.get() )

tooptions = [
    "jpg",
    "png",
    "jpeg",
    "ico",
    "txt",
    "pdf",
    "docs",
    "mp3",
    "wav",
    "mp4",
    "gif"
]
clickedto.set( "jpg" )
dropto = OptionMenu( Screen , clickedto , *tooptions )
dropto.place(relx=0.1,rely=0.76)
dropto.configure(bg='#00A6ED')

toBtn = Button(f1, padx = 5, pady = 3, bd = 5, fg = "black",
font = ('arial', 16, 'bold'), width = 10,
text = "click Me", bg = "#A8B4A5",
command = showto).grid(row = 3, column = 1)


labelto = Label( f1 , font = ('arial', 10, 'bold'), text = " ", bd = 5, anchor = "w" )
labelto.grid(row = 3, column = 2)













searchFile = Button(f1, padx = 5, pady = 3, bd = 5, fg = "black",
font = ('arial', 16, 'bold'), width = 10,
text = "Search file", bg = "#00A6ED",
command = OpenFile).grid(row = 1, column = 0)

# 2 Empty columns for blank spaces
emptylabel = Label(f1, font = ('arial', 16, 'bold'), text = "",width=10,  bd = 16, anchor = "w")
emptylabel.grid(row = 0, column = 3)
emptylabel.configure(bg='#A3BBAD')
emptylabel1 = Label(f1, font = ('arial', 16, 'bold'), text = "", bd = 16, anchor = "w")
emptylabel1.grid(row = 0, column = 4)
emptylabel1.configure(bg='#A3BBAD')    

# Column for file name display
addrReference = Entry(f1, font = ('arial', 16, 'bold'), textvariable = addr, justify = 'left')
addrReference.grid(row = 1, column = 1)

# Empty column for blank spaces
emptylabel2 = Label(f1, font = ('arial', 16, 'bold'), text = "", bd = 16, anchor = "w")
emptylabel2.grid(row = 2, column = 0)
emptylabel2.configure(bg='#A3BBAD')
emptylabel3 = Label(f1, font = ('arial', 16, 'bold'), text = "",width=10,  bd = 16, anchor = "w")
emptylabel3.grid(row = 3, column = 0)
emptylabel3.configure(bg='#A3BBAD')

processIntr = Button(f1, padx = 5, pady = 3, bd = 5, fg = "black",
font = ('arial', 16, 'bold'), width = 10,
text = "Convert", bg = "#A8B4A5",
command = runCode).grid(row = 0, column = 6)

# Column for show result
showResult = Label(f1, font = ('arial', 16, 'bold'), text = "Result :", bd = 16, anchor = "w")
showResult.grid(row = 4,  column = 4)
showResult.configure(bg='#A3BBAD')

# Column for show result output
resultMessage = Entry(f1, font = ('arial', 16, 'bold'), textvariable = Result, bd = 10, insertwidth = 4, width=30, bg = "white", justify = 'left')
resultMessage.grid(row = 4, columnspan= 2, column = 5)

# Empty column for blank spaces
emptylabel2 = Label(f1, font = ('arial', 16, 'bold'), text = "", bd = 16, anchor = "w")
emptylabel2.grid(row = 5, column = 0)
emptylabel2.configure(bg='#A3BBAD')

  


Screen.mainloop()